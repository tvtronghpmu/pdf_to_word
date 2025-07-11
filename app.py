import streamlit as st
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
import os
import io
import re # Thư viện xử lý biểu thức chính quy (Regular Expression)

# --- TẠO HÀM LÀM SẠCH ---
def sanitize_text_for_xml(text_string: str) -> str:
    """
    Hàm này loại bỏ các ký tự không tương thích với XML từ một chuỗi.
    Cụ thể là các ký tự điều khiển C0 và C1, ngoại trừ tab, xuống dòng và về đầu dòng.
    """
    if not isinstance(text_string, str):
        return ""
    # Sử dụng regex để tìm và thay thế các ký tự điều khiển không hợp lệ bằng chuỗi rỗng
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text_string)

# --- Tạo thư mục lưu file output ---
output_folder = "converted_docs"
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# --- HÀM XỬ LÝ CHÍNH ---

def convert_pdf_to_word(uploaded_file):
    """
    Hàm chính để chuyển đổi file PDF được tải lên sang file Word.
    Hàm sẽ xử lý cả văn bản thuần túy và hình ảnh (sử dụng OCR).
    """
    try:
        # Mở file PDF từ dữ liệu trong bộ nhớ
        pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        
        # Tạo một tài liệu Word mới
        doc = Document()

        # Lặp qua từng trang của file PDF
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # --- 1. Trích xuất và làm sạch văn bản thuần túy ---
            text = page.get_text("text")
            cleaned_text = sanitize_text_for_xml(text)
            if cleaned_text.strip():
                doc.add_paragraph(cleaned_text)

            # --- 2. Trích xuất hình ảnh và sử dụng OCR ---
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                image = Image.open(io.BytesIO(image_bytes))

                # Sử dụng Tesseract để nhận diện văn bản từ hình ảnh
                ocr_text = pytesseract.image_to_string(image, lang='vie+eng')
                # Làm sạch văn bản từ OCR
                cleaned_ocr_text = sanitize_text_for_xml(ocr_text)
                
                if cleaned_ocr_text.strip():
                    doc.add_paragraph("[Văn bản từ hình ảnh (OCR)]:")
                    doc.add_paragraph(cleaned_ocr_text)

            # Thêm dấu ngắt trang để giữ nguyên bố cục tương đối
            if page_num < len(pdf_document) - 1:
                doc.add_page_break()

        # Tạo đường dẫn lưu file output
        output_filename = os.path.splitext(uploaded_file.name)[0] + ".docx"
        output_path = os.path.join(output_folder, output_filename)
        
        # Lưu tài liệu Word
        doc.save(output_path)
        
        return output_path

    except Exception as e:
        # Hiển thị lỗi một cách chi tiết hơn để dễ gỡ rối
        st.error(f"Đã xảy ra lỗi trong quá trình xử lý: {e}")
        # Ghi log lỗi ra console của Streamlit để xem chi tiết
        print(f"Error details: {e}")
        return None

# --- GIAO DIỆN NGƯỜI DÙNG VỚI STREAMLIT ---

st.title("Chuyển đổi PDF sang Word (với OCR) 📄➡️📝")
st.write("Tải lên file PDF của bạn, ứng dụng sẽ chuyển đổi nó sang định dạng .docx.")

uploaded_file = st.file_uploader("Chọn một file PDF", type=["pdf"])

if st.button("Chuyển đổi"):
    if uploaded_file is not None:
        with st.spinner('Đang xử lý... Quá trình này có thể mất vài phút.'):
            result_path = convert_pdf_to_word(uploaded_file)
        
        if result_path:
            st.success("Chuyển đổi thành công!")
            st.info(f"File đã được lưu tại thư mục cục bộ: `{result_path}`")
            
            with open(result_path, "rb") as f:
                st.download_button(
                    label="Tải file Word (.docx)",
                    data=f,
                    file_name=os.path.basename(result_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.warning("Vui lòng tải lên một file PDF trước.")
